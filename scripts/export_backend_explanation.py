from pathlib import Path
from datetime import date

from docx import Document
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import cm
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, ListFlowable, ListItem

ROOT = Path(__file__).resolve().parents[1]
DOCS_DIR = ROOT / "docs"
DOCS_DIR.mkdir(parents=True, exist_ok=True)

DOCX_PATH = DOCS_DIR / "FitCoach_Backend_Explanation.docx"
PDF_PATH = DOCS_DIR / "FitCoach_Backend_Explanation.pdf"

TODAY = date.today().isoformat()

sections = [
    {
        "title": "FitCoach Backend Explanation (analysis_logic.py + app.py)",
        "paragraphs": [
            f"Prepared on {TODAY}",
            "This document explains the backend architecture so it can be presented clearly to teammates.",
            "The backend is split into two main responsibilities:",
        ],
        "bullets": [
            "analysis_logic.py: fitness analytics and deterministic calculations.",
            "app.py: API routing, intent/language handling, reply orchestration, and endpoint responses.",
        ],
    },
    {
        "title": "1) High-Level Architecture",
        "paragraphs": [
            "Frontend sends a request to POST /api/chat with: message, optional analytics, and optional language.",
            "app.py validates request schema, detects language/intent, and decides the response path.",
            "If analytics are present, app.py calls analysis_logic.py to compute metrics and include them in output.",
            "Finally, app.py builds the final reply (progress/food/nutrition/workout/fallback) and returns {reply, metrics}.",
        ],
    },
    {
        "title": "2) analysis_logic.py Detailed Walkthrough",
        "paragraphs": [
            "analysis_logic.py is designed as a pure logic module with testable functions.",
        ],
        "bullets": [
            "DEFAULT_TARGETS: default weekly change targets for weight/fat loss and weight/muscle gain.",
            "_to_float(value, fallback): safe numeric conversion to avoid crashes on malformed input.",
            "_clamp(value, low, high): keeps values within a bounded range.",
            "_compute_training_volume(workouts): computes weekly volume = sum(sets * reps * load).",
        ],
    },
    {
        "title": "2.1 weekly_progress_rate(current_week, previous_week)",
        "paragraphs": [
            "Purpose: compare current week performance versus previous week.",
            "Weighted score formula:",
            "Progress Rate (%) = 0.7 * Volume Change (%) + 0.3 * Adherence Change (%)",
            "Output includes status, current/previous volume, adherence %, and final progress_rate_pct.",
            "If previous week is missing, it returns status=insufficient_data.",
        ],
    },
    {
        "title": "2.2 goal_achievement_percentage(goal, profile, monthly_stats)",
        "paragraphs": [
            "Purpose: calculate goal completion percent by goal type.",
        ],
        "bullets": [
            "weight_loss / fat_loss: (start_weight - current_weight) / (start_weight - target_weight).",
            "weight_gain: (current_weight - start_weight) / (target_weight - start_weight).",
            "strength_gain: current strength increase / target strength increase.",
            "muscle_gain: weighted score = 70% strength component + 30% weight component.",
            "All outputs are clamped to 0-100%; invalid targets return None.",
        ],
    },
    {
        "title": "2.3 time_to_goal_estimation(goal, profile, weekly_trend_last_4)",
        "paragraphs": [
            "Purpose: estimate remaining weeks to goal based on recent trend average.",
            "Handles goal direction correctly:",
        ],
        "bullets": [
            "Loss goals require negative trend (weight decreasing).",
            "Gain goals require positive trend (weight increasing).",
            "If trend direction is wrong or zero, status becomes stalled_or_off_track.",
            "If already achieved, eta_weeks = 0 with status achieved.",
            "If insufficient trend data (<2), status insufficient_data.",
        ],
    },
    {
        "title": "2.4 calorie_adjustment_engine(...)",
        "paragraphs": [
            "Purpose: recommend calorie adjustment from actual vs target weekly weight change.",
        ],
        "bullets": [
            "Uses defaults from DEFAULT_TARGETS when target_weekly_change_kg is not provided.",
            "For fat/weight loss goals: slows/accelerates deficit depending on progress speed.",
            "For gain goals: increases or trims surplus based on gain speed.",
            "Adjustment is bounded by max_adjustment and enforced min_calories.",
            "Returns reason string to explain recommendation logic.",
        ],
    },
    {
        "title": "2.5 goal_tracking_function(...)",
        "paragraphs": [
            "Purpose: provide a high-level goal status for UI and coaching logic.",
        ],
        "bullets": [
            "ACHIEVED: progress_pct >= 100.",
            "STALLED: ETA logic reports stalled_or_off_track.",
            "IN_PROGRESS: default state otherwise.",
            "Returns status + goal type + progress % + ETA details.",
        ],
    },
    {
        "title": "3) app.py Detailed Walkthrough",
        "paragraphs": [
            "app.py is the orchestration layer that wraps all logic behind FastAPI endpoints.",
        ],
        "bullets": [
            "Defines request/response models with Pydantic (ChatRequest, ChatResponse, WelcomeResponse).",
            "Loads JSON knowledge sources at startup: responses, intents, foods, nutrition programs.",
            "Configures CORS for frontend access.",
            "Keeps helper functions for language handling, intent matching, and text normalization.",
        ],
    },
    {
        "title": "3.1 Language and Intent Routing",
        "paragraphs": [
            "Language selection strategy:",
        ],
        "bullets": [
            "If Arabic characters exist in message => Arabic response.",
            "Else, use explicit req.language if provided.",
            "Else default to English.",
            "Intent matching first tries pattern-based matching from conversation_intents.json.",
            "Then keyword checks route messages to progress, food, nutrition, or workout branches.",
        ],
    },
    {
        "title": "3.2 Food Reply Pipeline",
        "paragraphs": [
            "Food flow uses foods_data.json with robust matching and cleaning.",
        ],
        "bullets": [
            "_find_food_item: exact/substring normalized matching, then fuzzy fallback with SequenceMatcher.",
            "Arabic normalization removes diacritics and letter variants to improve matching quality.",
            "_build_food_reply: supports targeted questions (benefits, nutrients, suitable_for, alternatives).",
            "If no specific angle asked, returns concise summary with category, key nutrient, and note.",
        ],
    },
    {
        "title": "3.3 Nutrition Program Pipeline",
        "paragraphs": [
            "Nutrition flow uses nutrition_programs.json.",
        ],
        "bullets": [
            "Extracts weight from message (e.g., 80 kg).",
            "Detects goal hint (cutting or bulking).",
            "Filters candidate programs by goal and weight range.",
            "Builds bilingual response including calories, macro split, sample meal, and one tip.",
        ],
    },
    {
        "title": "3.4 Progress Reply Pipeline",
        "paragraphs": [
            "If message is progress-related and analytics payload exists, app.py builds a progress explanation.",
            "It formats template variables such as weekly change, remaining kg, ETA weeks, and calorie delta.",
            "Template source is responses.json (progress_analysis templates).",
        ],
    },
    {
        "title": "3.5 Core Decision Function: _build_chat_reply(req, metrics)",
        "paragraphs": [
            "Response priority order is important for predictability:",
        ],
        "bullets": [
            "1) Pattern-based intent reply (conversation_intents).",
            "2) Greeting response.",
            "3) Progress response (if intent + analytics).",
            "4) Food response.",
            "5) Nutrition response.",
            "6) Workout template response.",
            "7) Out-of-scope fallback.",
        ],
    },
    {
        "title": "3.6 API Endpoints",
        "bullets": [
            "GET /health: simple service status check.",
            "GET /api/chat/welcome: language-aware welcome message.",
            "POST /api/chat: main endpoint, computes metrics (if analytics), then returns final reply and metrics.",
        ],
    },
    {
        "title": "4) How app.py and analysis_logic.py Work Together",
        "paragraphs": [
            "The integration contract is simple:",
        ],
        "bullets": [
            "analysis_logic.py receives structured analytics and returns deterministic metric objects.",
            "app.py decides when to call those functions and inserts results into response payload.",
            "Chat reply generation can use these metrics for progress explanations.",
            "Frontend gets both narrative reply and machine-readable metrics for UI features.",
        ],
    },
    {
        "title": "5) Team-Friendly Summary",
        "bullets": [
            "analysis_logic.py = 'calculation brain' (testable fitness analytics).",
            "app.py = 'conversation orchestrator + API layer'.",
            "Current system is rule/data-driven (JSON + Python logic), not LLM-powered generation.",
            "This separation keeps logic maintainable, explainable, and easy to expand later.",
        ],
    },
]


def export_docx() -> None:
    doc = Document()
    for idx, section in enumerate(sections):
        if idx == 0:
            doc.add_heading(section["title"], level=0)
        else:
            doc.add_heading(section["title"], level=1)

        for para in section.get("paragraphs", []):
            doc.add_paragraph(para)

        for bullet in section.get("bullets", []):
            doc.add_paragraph(bullet, style="List Bullet")

    doc.save(DOCX_PATH)


def export_pdf() -> None:
    doc = SimpleDocTemplate(
        str(PDF_PATH),
        pagesize=A4,
        rightMargin=2 * cm,
        leftMargin=2 * cm,
        topMargin=2 * cm,
        bottomMargin=2 * cm,
    )

    styles = getSampleStyleSheet()
    heading_style = styles["Heading1"]
    normal_style = styles["BodyText"]
    normal_style.spaceAfter = 8

    title_style = ParagraphStyle(
        "TitleCustom",
        parent=styles["Title"],
        fontSize=18,
        leading=22,
        spaceAfter=12,
    )

    story = []
    for idx, section in enumerate(sections):
        if idx == 0:
            story.append(Paragraph(section["title"], title_style))
        else:
            story.append(Paragraph(section["title"], heading_style))

        for para in section.get("paragraphs", []):
            story.append(Paragraph(para, normal_style))

        bullets = section.get("bullets", [])
        if bullets:
            bullet_items = [
                ListItem(Paragraph(item, normal_style), leftIndent=10) for item in bullets
            ]
            story.append(ListFlowable(bullet_items, bulletType="bullet", leftIndent=12))
            story.append(Spacer(1, 8))

        story.append(Spacer(1, 6))

    doc.build(story)


if __name__ == "__main__":
    export_docx()
    export_pdf()
    print(f"Created: {DOCX_PATH}")
    print(f"Created: {PDF_PATH}")
